# -*- coding: utf-8 -*-
"""
GERADOR DE FOLHA DE ALTERAÇÕES
Extrai menções a um policial específico de boletins internos em PDF
e gera um documento Word (.docx) no formato institucional.

Uso interativo:
    python folha_alteracoes.py

Uso com argumentos:
    python folha_alteracoes.py --pasta "C:\\PDFs" --busca "NOME COMPLETO" --ano 2025 --semestre 1

Dependências:
    pip install pdfplumber python-docx
    (Opcional para PDFs escaneados: pip install pytesseract pdf2image + Tesseract OCR + Poppler)
"""

import os
import re
import sys
import argparse
from collections import defaultdict

import pdfplumber
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from pdf2image import convert_from_path
    import pytesseract

    TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    if os.path.exists(TESSERACT_PATH):
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

    POPPLER_PATH = None
    _candidatos_poppler = [
        r"C:\Users\levyg\AppData\Local\Microsoft\WinGet\Packages\oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe\poppler-25.07.0\Library\bin",
        r"C:\Program Files\poppler\Library\bin",
        r"C:\Program Files\poppler-24.08.0\Library\bin",
    ]
    for _p in _candidatos_poppler:
        if os.path.isdir(_p):
            POPPLER_PATH = _p
            break

    OCR_DISPONIVEL = True
except ImportError:
    OCR_DISPONIVEL = False
    POPPLER_PATH = None

MESES_ORDEM = [
    "JANEIRO", "FEVEREIRO", "MARÇO", "ABRIL", "MAIO", "JUNHO",
    "JULHO", "AGOSTO", "SETEMBRO", "OUTUBRO", "NOVEMBRO", "DEZEMBRO",
]

MESES_PARA_NUM = {}
for _i, _m in enumerate(MESES_ORDEM, 1):
    MESES_PARA_NUM[_m] = _i
MESES_PARA_NUM.update({
    "MARCO": 3,
    "JAN": 1, "FEV": 2, "MAR": 3, "ABR": 4, "MAI": 5, "JUN": 6,
    "JUL": 7, "AGO": 8, "SET": 9, "OUT": 10, "NOV": 11, "DEZ": 12,
})


def normalizar_mes(texto):
    """Normaliza nome de mês para número (1-12)."""
    t = texto.upper().strip()
    t = t.replace("Ç", "C").replace("ç", "c")
    for nome, num in MESES_PARA_NUM.items():
        nome_norm = nome.replace("Ç", "C")
        if t == nome_norm or t == nome:
            return num
    return None


def extrair_texto_pdf(caminho_pdf):
    """Extrai texto do PDF. Tenta pdfplumber; se vazio, tenta OCR."""
    texto_completo = ""
    tem_texto = False

    try:
        with pdfplumber.open(caminho_pdf) as pdf:
            for pagina in pdf.pages:
                texto = pagina.extract_text()
                if texto and texto.strip():
                    texto_completo += texto + "\n"
                    tem_texto = True
    except Exception as e:
        print(f"  [ERRO] Falha ao ler PDF: {e}")
        return None

    if tem_texto:
        return texto_completo

    if OCR_DISPONIVEL:
        print(f"[OCR]", end=" ", flush=True)
        try:
            kwargs = {"dpi": 300}
            if POPPLER_PATH:
                kwargs["poppler_path"] = POPPLER_PATH
            imagens = convert_from_path(caminho_pdf, **kwargs)
            for img in imagens:
                texto = pytesseract.image_to_string(img, lang="por")
                texto_completo += texto + "\n"
            if texto_completo.strip():
                return texto_completo
        except Exception as e:
            print(f"\n  [AVISO] OCR falhou: {e}")

    return None


def detectar_info_boletim(texto, nome_arquivo=None):
    """Extrai número do boletim, data, mês e ano do cabeçalho, página ou nome do arquivo."""
    info = {"numero": None, "data_str": None, "mes_num": None, "ano": None, "unidade": None}

    texto_upper = texto.upper()

    # Detectar unidade
    if "BATALHÃO ESPECIALIZADO" in texto_upper or "BEPE" in texto_upper:
        info["unidade"] = "BEPE"
    elif "BATALHÃO DE POLÍCIA DE CHOQUE" in texto_upper or "BPCHQ" in texto_upper:
        info["unidade"] = "BPChq"
    else:
        m = re.search(r"BATALHÃO\s+(?:DE\s+)?(.+?)(?:\n|$)", texto, re.IGNORECASE)
        if m:
            info["unidade"] = m.group(1).strip()[:30]

    # Detectar número do boletim
    padroes_numero = [
        r"BOLETIM\s+INTERNO\s*(?:DO\s+\w+)?\s*\n\s*N[.ºo°]*\s*(\d{2,3})",
        r"Boletim\s+Interno\s+n[.ºo°]\s*(\d{2,3})",
        r"N[.ºo°]\s*(\d{2,3})\s*$",
    ]
    for padrao in padroes_numero:
        m = re.search(padrao, texto, re.IGNORECASE | re.MULTILINE)
        if m:
            info["numero"] = m.group(1)
            break

    # Detectar data - padrão principal: "Rio de Janeiro, DD de MÊS de AAAA"
    m = re.search(
        r"Rio\s+de\s+Janeiro\s*[,.]?\s*(\d{1,2})\s+de\s+(\w+)\s+de\s+(\d{4})",
        texto, re.IGNORECASE,
    )
    if m:
        dia, mes_txt, ano = m.group(1), m.group(2), m.group(3)
        mes_num = normalizar_mes(mes_txt)
        if mes_num:
            info["mes_num"] = mes_num
            info["ano"] = int(ano)
            info["data_str"] = f"{dia} de {mes_txt.upper()} de {ano}"

    # Fallback data: do cabeçalho de página do boletim
    if not info["mes_num"]:
        m = re.search(
            r"BOL\.?\s*INT\.?\s*(?:N[ºo°]?\.?\s*)?\d+\s*(?:de\s+)?(\d{1,2})\s*de\s+(\w+)\s+de\s+(\d{4})",
            texto, re.IGNORECASE,
        )
        if m:
            dia, mes_txt, ano = m.group(1), m.group(2), m.group(3)
            mes_num = normalizar_mes(mes_txt)
            if mes_num:
                info["mes_num"] = mes_num
                info["ano"] = int(ano)
                info["data_str"] = f"{dia} de {mes_txt.upper()} de {ano}"

    # Fallback: extrair do nome do arquivo
    if nome_arquivo and (not info["mes_num"] or not info["numero"]):
        # Formato "BOL INT 037 08JUL2025"
        m = re.search(r"BOL\s*INT\s*(\d{2,3})\s+(\d{2})(\w{3})(\d{4})", nome_arquivo, re.IGNORECASE)
        if m:
            if not info["numero"]:
                info["numero"] = m.group(1)
            if not info["mes_num"]:
                dia, mes_abrev, ano = m.group(2), m.group(3).upper(), m.group(4)
                mes_num = normalizar_mes(mes_abrev)
                if mes_num:
                    info["mes_num"] = mes_num
                    info["ano"] = int(ano)
                    info["data_str"] = f"{dia} de {MESES_ORDEM[mes_num - 1]} de {ano}"
        else:
            # Formato "BOL INT 016 20.10.2025" ou "BOL INT 003 16.01.2025"
            m = re.search(r"BOL\s*INT\s*(\d{2,3})\s+(\d{2})[.\-](\d{2})[.\-](\d{2,4})", nome_arquivo, re.IGNORECASE)
            if m:
                if not info["numero"]:
                    info["numero"] = m.group(1)
                if not info["mes_num"]:
                    dia, mes_str, ano_str = m.group(2), m.group(3), m.group(4)
                    mes_num = int(mes_str)
                    ano = int(ano_str)
                    if ano < 100:
                        ano += 2000
                    if 1 <= mes_num <= 12:
                        info["mes_num"] = mes_num
                        info["ano"] = ano
                        info["data_str"] = f"{dia} de {MESES_ORDEM[mes_num - 1]} de {ano}"

    return info


def montar_referencia_boletim(info):
    """Monta a linha de referência do boletim."""
    partes = []
    if info.get("unidade"):
        partes.append(info["unidade"])
    partes.append("BOL. INT.")
    if info.get("numero"):
        partes.append(f"Nº {info['numero']}")
    if info.get("data_str"):
        partes.append(f"de {info['data_str']}")
    return " ".join(partes)


# Padrões de cabeçalho de página a remover na limpeza
_PADROES_CABECALHO = [
    re.compile(r"^BEPE\s*[–-]\s*Boletim\s+Interno", re.IGNORECASE),
    re.compile(r"^Bpchq\s+BOL\.\s+INT\.\s+N", re.IGNORECASE),
    re.compile(r"^SERGIO\s+WILLIAM", re.IGNORECASE),
    re.compile(r"^TEN\s+CEL\s+PM\s+RG", re.IGNORECASE),
    re.compile(r"^Página\s+\d+\s+de\s+\d+", re.IGNORECASE),
    re.compile(r"^S[ée]rgio\s+Wi", re.IGNORECASE),
    re.compile(r"^TEN\.\s*CEL", re.IGNORECASE),
    re.compile(r"^ID\.\s*\d+", re.IGNORECASE),
]


def dividir_em_entradas(texto):
    """Divide o texto do boletim em entradas individuais (itens numerados e seções)."""
    linhas = texto.split("\n")
    entradas = []
    entrada_atual = []
    titulo_atual = ""

    padrao_inicio = re.compile(
        r"^\s*(\d{1,3})\s*[-–]\s*[A-ZÀ-ÚÇa-zà-úç]|"
        r"^\s*(\d{1,3})\.\s+[A-ZÀ-ÚÇ]"
    )

    padrao_secao = re.compile(
        r"^\s*(?:[A-Z]\)\s+ALTERAÇ|[IVX]+\.\s+ASSUNTOS|"
        r"\d+[ªaº]\s*\.?\s*PARTE|"
        r"B\)\s+ALTERAÇ|A\)\s+ALTERAÇ)",
        re.IGNORECASE,
    )

    for linha in linhas:
        linha_strip = linha.strip()
        if padrao_inicio.match(linha_strip) or padrao_secao.match(linha_strip):
            if entrada_atual:
                entradas.append({
                    "titulo": titulo_atual,
                    "texto": "\n".join(entrada_atual),
                })
            entrada_atual = [linha]
            titulo_atual = linha_strip
        else:
            entrada_atual.append(linha)

    if entrada_atual:
        entradas.append({
            "titulo": titulo_atual,
            "texto": "\n".join(entrada_atual),
        })

    return entradas


def limpar_entrada(texto):
    """Remove cabeçalhos de página repetidos e linhas de paginação."""
    linhas = texto.split("\n")
    linhas_limpas = []
    for linha in linhas:
        l = linha.strip()
        if l == "-" or l == "|" or l == "—":
            continue
        if any(p.match(l) for p in _PADROES_CABECALHO):
            continue
        linhas_limpas.append(linha)

    while linhas_limpas and not linhas_limpas[0].strip():
        linhas_limpas.pop(0)
    while linhas_limpas and not linhas_limpas[-1].strip():
        linhas_limpas.pop()

    return "\n".join(linhas_limpas)


def buscar_pessoa(texto_completo, termo_busca):
    """Busca o termo (nome ou RG) no texto e retorna as entradas correspondentes."""
    resultados = []
    termo_upper = termo_busca.upper().strip()
    termo_sem_pontos = termo_upper.replace(".", "").replace("-", "").replace(" ", "")

    entradas = dividir_em_entradas(texto_completo)

    for entrada in entradas:
        texto_entrada = entrada["texto"].upper()
        texto_norm = texto_entrada.replace(".", "").replace("-", "").replace(" ", "")

        if termo_upper in texto_entrada or termo_sem_pontos in texto_norm:
            texto_limpo = limpar_entrada(entrada["texto"])
            if texto_limpo.strip():
                resultados.append(texto_limpo)

    return resultados


def processar_pasta(pasta_pdfs, termo_busca):
    """Processa todos os PDFs da pasta e retorna ocorrências organizadas por (ano, mês)."""
    ocorrencias = defaultdict(list)
    total_arquivos = 0
    total_ocorrencias = 0
    arquivos_sem_texto = []

    arquivos_pdf = []
    for raiz, _, arquivos in os.walk(pasta_pdfs):
        for arq in arquivos:
            if arq.lower().endswith(".pdf"):
                arquivos_pdf.append(os.path.join(raiz, arq))

    arquivos_pdf.sort()
    print(f"\n{'=' * 70}")
    print(f"  PROCESSANDO {len(arquivos_pdf)} ARQUIVO(S) PDF")
    print(f"  Buscando: {termo_busca}")
    print(f"{'=' * 70}\n")

    for caminho in arquivos_pdf:
        nome_arquivo = os.path.basename(caminho)
        total_arquivos += 1
        print(f"[{total_arquivos:3d}/{len(arquivos_pdf)}] {nome_arquivo}... ", end="", flush=True)

        texto = extrair_texto_pdf(caminho)
        if not texto:
            print("SEM TEXTO EXTRAÍVEL")
            arquivos_sem_texto.append(nome_arquivo)
            continue

        info = detectar_info_boletim(texto, nome_arquivo)
        if not info["mes_num"] or not info["ano"]:
            print("DATA NÃO DETECTADA")
            arquivos_sem_texto.append(nome_arquivo)
            continue

        resultados = buscar_pessoa(texto, termo_busca)

        if resultados:
            ref_boletim = montar_referencia_boletim(info)
            chave = (info["ano"], info["mes_num"])
            for trecho in resultados:
                ocorrencias[chave].append({
                    "referencia": ref_boletim,
                    "texto": trecho,
                })
            total_ocorrencias += len(resultados)
            print(f">>> {len(resultados)} OCORRÊNCIA(S) <<<")
        else:
            print("nenhuma ocorrência")

    print(f"\n{'=' * 70}")
    print(f"  RESUMO")
    print(f"  Arquivos processados: {total_arquivos}")
    print(f"  Ocorrências encontradas: {total_ocorrencias}")
    if arquivos_sem_texto:
        print(f"  Arquivos com problemas: {len(arquivos_sem_texto)}")
        for a in arquivos_sem_texto:
            print(f"    - {a}")
    print(f"{'=' * 70}\n")

    return ocorrencias, total_ocorrencias


def gerar_docx(ocorrencias, termo_busca, ano_ref, caminho_saida, semestre=None):
    """Gera o documento Word no formato Folha de Alterações."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    for _ in range(3):
        doc.add_paragraph("")

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("FOLHA DE ALTERAÇÕES")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Arial"

    doc.add_paragraph("")

    if semestre == 1:
        meses_gerar = MESES_ORDEM[:6]
        meses_nums = list(range(1, 7))
    elif semestre == 2:
        meses_gerar = MESES_ORDEM[6:]
        meses_nums = list(range(7, 13))
    else:
        meses_gerar = MESES_ORDEM
        meses_nums = list(range(1, 13))

    for mes_nome, mes_num in zip(meses_gerar, meses_nums):
        doc.add_paragraph("")

        p_mes = doc.add_paragraph()
        p_mes.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_mes = p_mes.add_run(mes_nome)
        run_mes.bold = True
        run_mes.font.size = Pt(12)
        run_mes.font.name = "Arial"

        doc.add_paragraph("")

        chave = (ano_ref, mes_num)
        entradas_mes = ocorrencias.get(chave, [])

        if not entradas_mes:
            p_sem = doc.add_paragraph()
            p_sem.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_sem = p_sem.add_run("SEM ALTERAÇÃO")
            run_sem.font.size = Pt(11)
            run_sem.font.name = "Arial"
        else:
            for entrada in entradas_mes:
                p_ref = doc.add_paragraph()
                p_ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run_ref = p_ref.add_run(entrada["referencia"])
                run_ref.bold = True
                run_ref.font.size = Pt(10)
                run_ref.font.name = "Arial"

                p_txt = doc.add_paragraph()
                p_txt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run_txt = p_txt.add_run(entrada["texto"])
                run_txt.font.size = Pt(10)
                run_txt.font.name = "Arial"

                doc.add_paragraph("")

    doc.add_paragraph("")
    doc.add_paragraph("")

    p_assin = doc.add_paragraph()
    p_assin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_assin = p_assin.add_run("_" * 40)
    run_assin.font.size = Pt(10)
    run_assin.font.name = "Arial"

    p_cmd = doc.add_paragraph()
    p_cmd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cmd = p_cmd.add_run("COMANDANTE DO BEPE")
    run_cmd.bold = True
    run_cmd.font.size = Pt(11)
    run_cmd.font.name = "Arial"

    doc.save(caminho_saida)
    print(f"  Documento gerado: {caminho_saida}")


def main():
    parser = argparse.ArgumentParser(
        description="Gerador de Folha de Alterações - Polícia Militar",
    )
    parser.add_argument("--pasta", help="Caminho da pasta com os PDFs")
    parser.add_argument("--busca", help="Nome completo ou RG do policial")
    parser.add_argument("--ano", type=int, help="Ano de referência (ex: 2025)")
    parser.add_argument("--semestre", type=int, choices=[0, 1, 2], default=0,
                        help="1=JAN-JUN, 2=JUL-DEZ, 0=ano completo")
    parser.add_argument("--saida", help="Caminho do arquivo de saída (.docx)")
    args = parser.parse_args()

    print("\n" + "=" * 70)
    print("  GERADOR DE FOLHA DE ALTERAÇÕES")
    print("  Polícia Militar - Extração de Boletins Internos")
    if OCR_DISPONIVEL:
        print("  [OCR disponível para PDFs escaneados]")
    else:
        print("  [OCR não disponível - apenas PDFs com texto]")
    print("=" * 70)

    pasta = args.pasta or input("\n  Caminho da pasta com os PDFs: ").strip().strip('"')
    if not os.path.isdir(pasta):
        print(f"  [ERRO] Pasta não encontrada: {pasta}")
        sys.exit(1)

    termo = args.busca or input("  Nome completo ou RG do policial: ").strip()
    if not termo:
        print("  [ERRO] Informe o nome ou RG.")
        sys.exit(1)

    if args.ano:
        ano_ref = args.ano
    else:
        ano_str = input("  Ano de referência (ex: 2025): ").strip()
        try:
            ano_ref = int(ano_str)
        except ValueError:
            print("  [ERRO] Ano inválido.")
            sys.exit(1)

    semestre = args.semestre
    if not args.pasta:
        sem_str = input("  Semestre (1=JAN-JUN, 2=JUL-DEZ, 0=ano completo) [0]: ").strip()
        try:
            semestre = int(sem_str) if sem_str else 0
        except ValueError:
            semestre = 0
        if semestre not in (0, 1, 2):
            semestre = 0

    if args.saida:
        caminho_saida = args.saida
    else:
        nome_saida = re.sub(r'[^\w\s\-]', '', termo).strip().replace(' ', '_')
        sufixo = {0: "ANO_COMPLETO", 1: "1o_SEMESTRE", 2: "2o_SEMESTRE"}[semestre]
        arquivo_saida = f"FOLHA_ALTERACOES_{nome_saida}_{sufixo}_{ano_ref}.docx"
        caminho_saida = os.path.join(os.getcwd(), arquivo_saida)

    ocorrencias, total = processar_pasta(pasta, termo)

    if total == 0:
        print(f"  Nenhuma ocorrência encontrada para '{termo}'.")
        print(f"  A Folha de Alterações será gerada com 'SEM ALTERAÇÃO' em todos os meses.\n")

    gerar_docx(ocorrencias, termo, ano_ref, caminho_saida, semestre if semestre else None)

    print(f"\n  Arquivo salvo em: {caminho_saida}")
    print(f"  Total de ocorrências: {total}")
    print("=" * 70)


if __name__ == "__main__":
    main()
